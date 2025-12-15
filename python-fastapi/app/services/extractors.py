from pathlib import Path
from typing import List

def _clean(s: str) -> str:
    # lightweight cleanup to avoid giant whitespace
    return "\n".join(line.rstrip() for line in s.splitlines()).strip()

def extract_text_from_file(path: Path) -> str:
    """
    Extracts text from:
      - .txt
      - .pdf
      - .docx
      - .csv
      - .xlsx
    Returns extracted text (possibly large). You can chunk later.
    """
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return _clean(path.read_text(encoding="utf-8", errors="ignore"))

    if suffix == ".pdf":
        return _clean(_extract_pdf(path))

    if suffix == ".docx":
        return _clean(_extract_docx(path))

    if suffix == ".csv":
        return _clean(_extract_csv(path))

    if suffix == ".xlsx":
        return _clean(_extract_xlsx(path))

    return f"[Unsupported file type: {suffix}]"


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: List[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        parts.append(f"--- Page {i+1} ---\n{text}")
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # tables (optional but often useful in syllabi)
    for t_i, table in enumerate(doc.tables, start=1):
        parts.append(f"\n--- Table {t_i} ---")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    import pandas as pd

    df = pd.read_csv(path, dtype=str, keep_default_na=False, encoding_errors="ignore")
    # Turn table into readable text
    return df.to_csv(index=False)


def _extract_xlsx(path: Path) -> str:
    import pandas as pd

    xls = pd.ExcelFile(path)
    parts: List[str] = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str, keep_default_na=False)
        parts.append(f"--- Sheet: {sheet_name} ---")
        parts.append(df.to_csv(index=False))
    return "\n\n".join(parts)
